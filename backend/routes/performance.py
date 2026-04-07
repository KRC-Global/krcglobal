"""
실적관리 (준공증명서) API
"""
from flask import Blueprint, jsonify, request, current_app, send_file
from models import db, PerformanceRecord, ConsultingProject
from routes.auth import token_required
from utils.file_naming import make_overseas_tech_filename, make_overseas_tech_disk_filename
from datetime import datetime
import os
import io
import tempfile
from werkzeug.utils import secure_filename

performance_bp = Blueprint('performance', __name__)

ALLOWED_EXTENSIONS = {'pdf', 'zip'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_project_date(date_str, is_end_date=False):
    """
    프로젝트 날짜 문자열을 date 객체로 변환
    지원 형식:
    - 'YYYY-MM-DD' (예: '1972-10-01')
    - 'YY-MM' (예: '72-10')
    - 'YYYY-MM' (예: '1972-10')
    """
    import calendar

    if not date_str:
        return None

    parts = str(date_str).split('-')

    if len(parts) == 3:
        # YYYY-MM-DD 형식
        year = int(parts[0])
        month = int(parts[1])
        day = int(parts[2])
        return datetime(year, month, day).date()

    elif len(parts) == 2:
        year_str = parts[0]
        month = int(parts[1])

        # 연도 파싱
        if len(year_str) == 4:
            year = int(year_str)
        elif len(year_str) == 2:
            year_num = int(year_str)
            # 70년 이상은 1900년대, 미만은 2000년대
            year = 1900 + year_num if year_num >= 70 else 2000 + year_num
        else:
            return None

        if is_end_date:
            # 준공일은 해당 월의 마지막 날
            last_day = calendar.monthrange(year, month)[1]
            return datetime(year, month, last_day).date()
        else:
            # 착수일은 해당 월의 1일
            return datetime(year, month, 1).date()

    return None


CONTINENT_COUNTRIES = {
    'asia': [
        '네팔', '라오스', '말레이시아', '몽골', '미얀마', '방글라데시', '베트남',
        '브루나이', '스리랑카', '아프가니스탄', '인도', '인도네시아', '중국',
        '캄보디아', '태국', '파키스탄', '필리핀', '이란', '키르기스스탄',
        '타지키스탄', '우즈베키스탄', '카자흐스탄', '투르크메니스탄',
    ],
    'africa': [
        'DR콩고', '가나', '말라위', '앙골라', '에티오피아', '우간다',
        '케냐', '코트디브아르', '탄자니아', '알제리', '모잠비크', '르완다',
        '세네갈', '카메룬', '나이지리아', '마다가스카르', '잠비아', '짐바브웨',
    ],
    'latinamerica': [
        '과테말라', '볼리비아', '아르헨티나', '엘살바도르', '파라과이',
        '페루', '콜롬비아', '에콰도르', '니카라과', '온두라스', '도미니카공화국',
    ],
    'oceania': [
        '키리바시', '피지', '파푸아뉴기니', '솔로몬제도', '사모아',
    ],
    'cis': [
        '러시아', '우즈베키스탄', '카자흐스탄', '키르기스스탄',
        '타지키스탄', '투르크메니스탄',
    ],
}


@performance_bp.route('', methods=['GET'])
@token_required
def get_records(current_user):
    """실적 목록 조회"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)

        # 필터링
        country = request.args.get('country')
        project_type = request.args.get('projectType')
        year = request.args.get('year', type=int)
        search = request.args.get('search')
        continent = request.args.get('continent')
        amount_range = request.args.get('amountRange')
        recent_years = request.args.get('recentYears', type=int)

        query = PerformanceRecord.query

        if country:
            query = query.filter(PerformanceRecord.country == country)
        if project_type:
            query = query.filter(PerformanceRecord.project_type == project_type)
        if year:
            query = query.filter(db.extract('year', PerformanceRecord.end_date) == year)
        if search:
            query = query.filter(
                db.or_(
                    PerformanceRecord.title.ilike(f'%{search}%'),
                    PerformanceRecord.client.ilike(f'%{search}%'),
                    PerformanceRecord.country.ilike(f'%{search}%')
                )
            )

        # 대륙별 필터
        if continent and continent in CONTINENT_COUNTRIES:
            countries_in_continent = CONTINENT_COUNTRIES[continent]
            query = query.filter(PerformanceRecord.country.in_(countries_in_continent))

        # 용역비 범위 필터 (US$ 백만 기준)
        if amount_range:
            if amount_range == 'under1':
                query = query.filter(PerformanceRecord.contract_amount_usd < 1)
            elif amount_range == '1to5':
                query = query.filter(
                    PerformanceRecord.contract_amount_usd >= 1,
                    PerformanceRecord.contract_amount_usd < 5
                )
            elif amount_range == '5to10':
                query = query.filter(
                    PerformanceRecord.contract_amount_usd >= 5,
                    PerformanceRecord.contract_amount_usd < 10
                )
            elif amount_range == 'over10':
                query = query.filter(PerformanceRecord.contract_amount_usd >= 10)

        # 최근 N년 이내 준공사업 필터
        if recent_years and recent_years > 0:
            from datetime import timedelta
            cutoff_date = datetime.utcnow().date() - timedelta(days=recent_years * 365)
            query = query.filter(PerformanceRecord.end_date >= cutoff_date)

        # 정렬 (준공일 기준 최신순)
        query = query.order_by(PerformanceRecord.end_date.desc(), PerformanceRecord.id.desc())

        # 페이지네이션
        pagination = query.paginate(page=page, per_page=per_page, error_out=False)

        return jsonify({
            'success': True,
            'data': [r.to_dict() for r in pagination.items],
            'total': pagination.total,
            'pages': pagination.pages,
            'currentPage': page,
            'perPage': per_page
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/<int:id>', methods=['GET'])
@token_required
def get_record(current_user, id):
    """실적 상세 조회"""
    try:
        record = PerformanceRecord.query.get_or_404(id)
        return jsonify({
            'success': True,
            'data': record.to_dict()
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('', methods=['POST'])
@token_required
def create_record(current_user):
    """실적 등록"""
    try:
        # 프로젝트 ID 확인
        consulting_project_id = request.form.get('consultingProjectId', type=int)

        if not consulting_project_id:
            return jsonify({'success': False, 'message': '사업을 선택해주세요.'}), 400

        # 프로젝트 존재 확인
        project = ConsultingProject.query.get(consulting_project_id)
        if not project:
            return jsonify({'success': False, 'message': '프로젝트를 찾을 수 없습니다.'}), 404

        # 프로젝트 정보로 자동 채우기
        record = PerformanceRecord(
            consulting_project_id=consulting_project_id,
            title=project.title_kr,
            country=project.country,
            client=project.client,
            funding_source=project.funding_source,
            project_type=request.form.get('projectType'),
            consortium_info=request.form.get('consortiumInfo'),
            description=request.form.get('description'),
            achievements=request.form.get('achievements'),
            remarks=request.form.get('remarks'),
            created_by=current_user.id
        )

        # 프로젝트의 계약금액과 공사지분 자동 설정
        if project.budget:
            record.contract_amount = project.budget
        if project.krc_budget:
            record.krc_amount = project.krc_budget
        # US$ 금액 자동 설정
        if project.budget_usd:
            record.contract_amount_usd = project.budget_usd
        if project.krc_budget_usd:
            record.krc_amount_usd = project.krc_budget_usd

        # 프로젝트의 착공일, 준공일 자동 설정
        if project.start_date:
            try:
                record.start_date = parse_project_date(project.start_date)
            except:
                pass

        if project.end_date:
            try:
                record.end_date = parse_project_date(project.end_date, is_end_date=True)
            except:
                pass

        # 날짜 처리 (수동 입력값이 있으면 프로젝트 기본값 덮어쓰기)
        contract_date = request.form.get('contractDate')
        if contract_date:
            record.contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()

        start_date = request.form.get('startDate')
        if start_date:
            record.start_date = datetime.strptime(start_date, '%Y-%m-%d').date()

        end_date = request.form.get('endDate')
        if end_date:
            record.end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

        # 파일 처리 (준공증명서)
        if 'file' in request.files:
            file = request.files['file']
            if file and file.filename and allowed_file(file.filename):
                original = secure_filename(file.filename)
                ext = original.rsplit('.', 1)[1].lower() if '.' in original else 'pdf'
                _fallback_year = record.contract_date.year if record.contract_date else None
                unique_filename = make_overseas_tech_disk_filename('준공증명서', ext, project, record.title, _fallback_year)

                upload_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'performance')
                os.makedirs(upload_dir, exist_ok=True)

                file_path = os.path.join(upload_dir, unique_filename)
                file.save(file_path)

                record.file_name = make_overseas_tech_filename('준공증명서', ext, project, record.title, _fallback_year)
                record.file_path = unique_filename  # 파일명만 저장 (플랫폼 독립적)
                record.file_size = os.path.getsize(file_path)

        db.session.add(record)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '실적이 등록되었습니다.',
            'data': record.to_dict()
        }), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/<int:id>', methods=['PUT'])
@token_required
def update_record(current_user, id):
    """실적 수정"""
    try:
        record = PerformanceRecord.query.get_or_404(id)

        # 프로젝트 변경 확인
        consulting_project_id = request.form.get('consultingProjectId', type=int)
        if consulting_project_id and consulting_project_id != record.consulting_project_id:
            # 프로젝트가 변경된 경우, 프로젝트 정보로 자동 업데이트
            project = ConsultingProject.query.get(consulting_project_id)
            if not project:
                return jsonify({'success': False, 'message': '프로젝트를 찾을 수 없습니다.'}), 404

            record.consulting_project_id = consulting_project_id
            record.title = project.title_kr
            record.country = project.country
            record.client = project.client
            record.funding_source = project.funding_source

            # 프로젝트의 계약금액과 공사지분 자동 설정
            if project.budget:
                record.contract_amount = project.budget
            if project.krc_budget:
                record.krc_amount = project.krc_budget
            # US$ 금액 자동 설정
            if project.budget_usd:
                record.contract_amount_usd = project.budget_usd
            if project.krc_budget_usd:
                record.krc_amount_usd = project.krc_budget_usd

            # 프로젝트의 착공일, 준공일 자동 설정
            if project.start_date:
                try:
                    record.start_date = parse_project_date(project.start_date)
                except:
                    pass

            if project.end_date:
                try:
                    record.end_date = parse_project_date(project.end_date, is_end_date=True)
                except:
                    pass

        # Form data 처리 (수동 수정값이 있으면 반영)
        if request.form.get('projectType'):
            record.project_type = request.form.get('projectType')
        if 'consortiumInfo' in request.form:
            record.consortium_info = request.form.get('consortiumInfo')
        if request.form.get('description'):
            record.description = request.form.get('description')
        if request.form.get('achievements'):
            record.achievements = request.form.get('achievements')
        if request.form.get('remarks'):
            record.remarks = request.form.get('remarks')

        # 날짜 처리
        contract_date = request.form.get('contractDate')
        if contract_date:
            record.contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()

        start_date = request.form.get('startDate')
        if start_date:
            record.start_date = datetime.strptime(start_date, '%Y-%m-%d').date()

        end_date = request.form.get('endDate')
        if end_date:
            record.end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

        # 새 파일 업로드
        if 'file' in request.files:
            file = request.files['file']
            if file and file.filename and allowed_file(file.filename):
                # 기존 파일 삭제
                if record.file_path and os.path.exists(record.file_path):
                    os.remove(record.file_path)

                original = secure_filename(file.filename)
                ext = original.rsplit('.', 1)[1].lower() if '.' in original else 'pdf'
                _upd_proj = ConsultingProject.query.get(record.consulting_project_id) if record.consulting_project_id else None
                _upd_year = record.contract_date.year if record.contract_date else None
                unique_filename = make_overseas_tech_disk_filename('준공증명서', ext, _upd_proj, record.title, _upd_year)

                upload_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'performance')
                os.makedirs(upload_dir, exist_ok=True)

                file_path = os.path.join(upload_dir, unique_filename)
                file.save(file_path)

                record.file_name = make_overseas_tech_filename('준공증명서', ext, _upd_proj, record.title, _upd_year)
                record.file_path = unique_filename  # 파일명만 저장 (플랫폼 독립적)
                record.file_size = os.path.getsize(file_path)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': '실적이 수정되었습니다.',
            'data': record.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/<int:id>', methods=['DELETE'])
@token_required
def delete_record(current_user, id):
    """실적 삭제"""
    try:
        record = PerformanceRecord.query.get_or_404(id)

        # 파일 삭제
        if record.file_path and os.path.exists(record.file_path):
            os.remove(record.file_path)

        db.session.delete(record)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '실적이 삭제되었습니다.'
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/<int:id>/download', methods=['GET'])
@token_required
def download_file(current_user, id):
    """준공증명서 파일 다운로드"""
    try:
        record = PerformanceRecord.query.get_or_404(id)

        if not record.file_path:
            return jsonify({'success': False, 'message': '파일 경로가 없습니다.'}), 404

        # 파일 경로 생성 (파일명만 저장되어 있으므로 upload_dir과 결합)
        upload_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'performance')
        file_path = os.path.join(upload_dir, record.file_path)

        if not os.path.exists(file_path):
            return jsonify({'success': False, 'message': '파일을 찾을 수 없습니다.'}), 404

        # 파일명 생성: 사업연도_사업명_준공증명서.ext
        _dl_proj = ConsultingProject.query.get(record.consulting_project_id) if record.consulting_project_id else None
        _dl_year = record.contract_date.year if record.contract_date else None

        # 실제 파일 확장자 추출
        file_ext = os.path.splitext(record.file_path)[1]  # .pdf, .zip 등
        if not file_ext:
            file_ext = '.pdf'  # 확장자가 없으면 기본값

        download_name = make_overseas_tech_filename('준공증명서', file_ext, _dl_proj, record.title, _dl_year)

        return send_file(
            file_path,
            as_attachment=True,
            download_name=download_name
        )

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/<int:id>/preview', methods=['GET'])
def preview_file(id):
    """준공증명서 파일 미리보기"""
    try:
        # 쿼리 파라미터 또는 헤더에서 토큰 가져오기
        token = request.args.get('token') or request.headers.get('Authorization', '').replace('Bearer ', '')

        if not token:
            return jsonify({'success': False, 'message': '인증 토큰이 필요합니다.'}), 401

        # 토큰 검증
        from routes.auth import verify_token
        user = verify_token(token)
        if not user:
            return jsonify({'success': False, 'message': '유효하지 않은 토큰입니다.'}), 401

        record = PerformanceRecord.query.get_or_404(id)

        if not record.file_path:
            return jsonify({'success': False, 'message': '파일 경로가 없습니다.'}), 404

        # 파일 경로 생성 (파일명만 저장되어 있으므로 upload_dir과 결합)
        upload_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'performance')
        file_path = os.path.join(upload_dir, record.file_path)

        if not os.path.exists(file_path):
            return jsonify({'success': False, 'message': '파일을 찾을 수 없습니다.'}), 404

        # PDF 파일을 브라우저에서 직접 표시 (미리보기 모드)
        return send_file(
            file_path,
            mimetype='application/pdf',
            as_attachment=False  # 미리보기 모드
        )

    except Exception as e:
        return jsonify({'success': False, 'message': f'미리보기 실패: {str(e)}'}), 500


@performance_bp.route('/stats', methods=['GET'])
@token_required
def get_stats(current_user):
    """실적 통계"""
    try:
        total = PerformanceRecord.query.count()

        # 총 계약금액 (KRW 백만원)
        total_amount = db.session.query(
            db.func.sum(PerformanceRecord.contract_amount)
        ).scalar() or 0

        # 총 공사지분 (KRW 백만원)
        total_krc = db.session.query(
            db.func.sum(PerformanceRecord.krc_amount)
        ).scalar() or 0

        # 총 계약금액 (US$ 백만)
        total_amount_usd = db.session.query(
            db.func.sum(PerformanceRecord.contract_amount_usd)
        ).scalar() or 0

        # 총 공사지분 (US$ 백만)
        total_krc_usd = db.session.query(
            db.func.sum(PerformanceRecord.krc_amount_usd)
        ).scalar() or 0

        # 국가별 실적 수
        by_country = db.session.query(
            PerformanceRecord.country,
            db.func.count(PerformanceRecord.id)
        ).group_by(PerformanceRecord.country).order_by(
            db.func.count(PerformanceRecord.id).desc()
        ).limit(10).all()

        return jsonify({
            'success': True,
            'data': {
                'total': total,
                'totalAmount': float(total_amount),
                'totalKrcAmount': float(total_krc),
                'totalAmountUsd': round(float(total_amount_usd), 2),
                'totalKrcAmountUsd': round(float(total_krc_usd), 2),
                'byCountry': [{'country': c[0], 'count': c[1]} for c in by_country if c[0]]
            }
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/countries', methods=['GET'])
@token_required
def get_countries(current_user):
    """실적 국가 목록"""
    try:
        countries = db.session.query(PerformanceRecord.country).distinct().filter(
            PerformanceRecord.country.isnot(None)
        ).order_by(PerformanceRecord.country).all()

        return jsonify({
            'success': True,
            'data': [c[0] for c in countries if c[0]]
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/project-types', methods=['GET'])
@token_required
def get_project_types(current_user):
    """실적 사업유형 목록"""
    try:
        types = db.session.query(PerformanceRecord.project_type).distinct().filter(
            PerformanceRecord.project_type.isnot(None)
        ).order_by(PerformanceRecord.project_type).all()

        return jsonify({
            'success': True,
            'data': [t[0] for t in types if t[0]]
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@performance_bp.route('/projects', methods=['GET'])
@token_required
def get_projects_for_selection(current_user):
    """실적 등록용 프로젝트 목록 조회 (준공년도 내림차순)"""
    try:
        # 준공년도(end_date) 기준 내림차순 정렬
        projects = ConsultingProject.query.order_by(
            ConsultingProject.end_date.desc().nullslast(),
            ConsultingProject.title_kr.asc()
        ).all()

        return jsonify({
            'success': True,
            'data': [p.to_dict() for p in projects]
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'프로젝트 목록 조회 실패: {str(e)}'}), 500


@performance_bp.route('/export-experience', methods=['POST'])
@token_required
def export_experience(current_user):
    """선택된 실적을 Consultant's Experience DOCX로 내보내기"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'success': False, 'message': '선택된 실적이 없습니다.'}), 400

        ids = data['ids']
        if not ids:
            return jsonify({'success': False, 'message': '선택된 실적이 없습니다.'}), 400

        # 선택된 실적 조회 (ID 순서 유지)
        records = PerformanceRecord.query.filter(
            PerformanceRecord.id.in_(ids)
        ).all()

        # ID 순서대로 정렬
        record_map = {r.id: r for r in records}
        ordered_records = [record_map[id] for id in ids if id in record_map]

        if not ordered_records:
            return jsonify({'success': False, 'message': '유효한 실적이 없습니다.'}), 404

        # DOCX 생성
        doc = Document()

        # 페이지 설정 (A4 세로, 좁은 마진)
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.08)
        section.right_margin = Cm(1.08)
        section.top_margin = Cm(1.44)
        section.bottom_margin = Cm(1.44)

        # 기본 폰트 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(9)

        # 제목 1: CONSULTANT'S EXPERIENCE
        title1 = doc.add_paragraph()
        title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = title1.add_run("CONSULTANT'S EXPERIENCE")
        run1.bold = True
        run1.font.name = 'Cambria'
        run1.font.size = Pt(12)

        # 제목 2: KRC
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = title2.add_run('KOREA RURAL COMMUNITY CORPORATION (KRC)')
        run2.bold = True
        run2.font.name = 'Cambria'
        run2.font.size = Pt(10)

        # 빈 줄
        doc.add_paragraph()

        # 테이블 생성 (헤더 + 데이터 행)
        num_rows = 1 + len(ordered_records)
        table = doc.add_table(rows=num_rows, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 컬럼 너비 설정
        col_widths = [Cm(0.8), Cm(2.5), Cm(8.0), Cm(3.0), Cm(3.0), Cm(2.5)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        # 헤더 행
        headers = [
            'No.',
            'Duration',
            'Assignment Name &\nBrief Description of Main\nDeliverables/Outputs',
            'Name of Client &\nCountry of\nAssignment',
            'Approx. Contract\nValue / Amount\nPaid to our Firm',
            'Role on the\nAssignment'
        ]

        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.bold = True
            run.font.name = 'Cambria'
            run.font.size = Pt(8)

            # 셀 세로 가운데 정렬
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = tcPr.makeelement(qn('w:vAlign'), {qn('w:val'): 'center'})
            tcPr.append(vAlign)

            # 셀 배경색 (연한 회색)
            shading = tcPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'D9E2F3'
            })
            tcPr.append(shading)

        # 데이터 행
        for row_idx, record in enumerate(ordered_records):
            row = table.rows[row_idx + 1]

            # 연결된 프로젝트 조회
            cp = record.consulting_project

            # Col 0: No. (빈칸 - 사용자가 직접 기입)
            cell0 = row.cells[0]
            cell0.text = ''
            cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Col 1: Duration (MMM. YYYY 형식)
            cell1 = row.cells[1]
            cell1.text = ''
            duration_text = _format_duration(record, cp)
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run(duration_text)
            run1.font.name = 'Cambria'
            run1.font.size = Pt(9)

            # Col 2: Assignment Name + Brief Description + Deliverables/Outputs
            cell2 = row.cells[2]
            cell2.text = ''
            _fill_assignment_cell(cell2, record, cp)

            # Col 3: Client / Country
            cell3 = row.cells[3]
            cell3.text = ''
            p3 = cell3.paragraphs[0]
            client_text = record.client or (cp.client if cp else '') or ''
            country_text = record.country or (cp.country if cp else '') or ''
            run3 = p3.add_run(f"{client_text}\n{country_text}")
            run3.font.name = 'Cambria'
            run3.font.size = Pt(9)

            # Col 4: Contract Value / Amount Paid
            cell4 = row.cells[4]
            cell4.text = ''
            p4 = cell4.paragraphs[0]
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contract_val = _format_usd_million(record.contract_amount_usd)
            krc_val = _format_usd_million(record.krc_amount_usd)
            run4 = p4.add_run(f"{contract_val}\n{krc_val}")
            run4.font.name = 'Cambria'
            run4.font.size = Pt(9)

            # Col 5: Role (auto-determine from consortium data)
            cell5 = row.cells[5]
            cell5.text = ''
            p5 = cell5.paragraphs[0]
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            role_text = _determine_role(record, cp)
            run5 = p5.add_run(role_text)
            run5.font.name = 'Cambria'
            run5.font.size = Pt(9)

        # 임시 파일로 저장 후 전송
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(tmp.name)
        tmp.close()

        return send_file(
            tmp.name,
            as_attachment=True,
            download_name='Consultants_Experience_KRC.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ImportError:
        return jsonify({'success': False, 'message': 'python-docx 라이브러리가 설치되지 않았습니다.'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': f'DOCX 생성 실패: {str(e)}'}), 500


def _format_duration(record, cp):
    """기간을 MMM. YYYY ~ MMM. YYYY 형식으로 변환"""
    import calendar

    def date_to_mmm_yyyy(d):
        """date 객체 또는 문자열을 MMM. YYYY 형식으로 변환"""
        if d is None:
            return ''

        if hasattr(d, 'strftime'):
            # date/datetime 객체
            return d.strftime('%b. %Y')

        # 문자열 파싱 (YY-MM 또는 YYYY-MM 또는 YYYY-MM-DD)
        d_str = str(d)
        parts = d_str.split('-')
        if len(parts) >= 2:
            year_str = parts[0]
            month = int(parts[1])

            if len(year_str) == 2:
                year_num = int(year_str)
                year = 1900 + year_num if year_num >= 70 else 2000 + year_num
            else:
                year = int(year_str)

            month_abbr = calendar.month_abbr[month]
            return f"{month_abbr}. {year}"

        return str(d)

    start = record.start_date
    end = record.end_date

    if not start and cp:
        start = cp.start_date
    if not end and cp:
        end = cp.end_date

    start_str = date_to_mmm_yyyy(start)
    end_str = date_to_mmm_yyyy(end)

    if start_str and end_str:
        return f"{start_str} ~\n{end_str}"
    elif start_str:
        return f"{start_str} ~"
    elif end_str:
        return f"~ {end_str}"
    return ''


def _determine_role(record, cp):
    """컨소시엄 데이터로 Role 자동 결정"""
    # record.project_type이 명시적으로 설정되어 있으면 우선 사용
    if record.project_type:
        return record.project_type

    if not cp:
        return ''

    lead = cp.lead_company or ''
    jv_partners = [jv for jv in [cp.jv1, cp.jv2, cp.jv3, cp.jv4, cp.jv5] if jv]
    has_jv = len(jv_partners) > 0

    # KRC가 주관사인지 확인
    krc_keywords = ['한국농어촌공사', 'KRC', 'Korea Rural Community', 'KRCC']
    is_krc_lead = any(kw.lower() in lead.lower() for kw in krc_keywords) if lead else False

    # KRC가 JV 파트너에 있는지 확인
    is_krc_jv = False
    for jv in jv_partners:
        if any(kw.lower() in jv.lower() for kw in krc_keywords):
            is_krc_jv = True
            break

    if is_krc_lead and has_jv:
        return 'Lead Firm in JV'
    elif is_krc_lead and not has_jv:
        return 'Sole Consultant'
    elif is_krc_jv:
        return 'Sub-contract'
    elif has_jv:
        # JV가 있지만 KRC 매칭 안됨 → lead가 비어있으면 Sole
        if not lead:
            return 'Sole Consultant' if not has_jv else 'Lead Firm in JV'
        return 'Sub-contract'

    return 'Sole Consultant'


def _build_consortium_text(record, cp):
    """컨소시엄 텍스트 생성 (프로젝트 데이터 우선, fallback: record.consortium_info)"""
    if cp:
        parts = []
        if cp.lead_company:
            parts.append(f"Lead: {cp.lead_company}")
        jv_list = []
        for i, jv in enumerate([cp.jv1, cp.jv2, cp.jv3, cp.jv4, cp.jv5], 1):
            if jv:
                jv_list.append(f"JV{i}: {jv}")
        parts.extend(jv_list)
        if parts:
            return ' / '.join(parts)

    return record.consortium_info or ''


def _fill_assignment_cell(cell, record, cp):
    """Assignment Name + Brief Description + Deliverables/Outputs 셀 작성"""
    from docx.shared import Pt

    # 기존 빈 paragraphs 삭제 후 작성
    p = cell.paragraphs[0]

    # Assignment Name (볼드)
    run_label = p.add_run('Assignment Name')
    run_label.bold = True
    run_label.font.name = 'Cambria'
    run_label.font.size = Pt(9)

    # 영문 사업명 구성
    title_en = (cp.title_en if cp and cp.title_en else None) or record.title
    funding = record.funding_source or (cp.funding_source if cp else None)

    if funding:
        assignment_name = f"[{funding} Funded] {title_en}"
    else:
        assignment_name = title_en

    # 새 줄에 사업명 (볼드)
    p2 = cell.add_paragraph()
    run_name = p2.add_run(assignment_name)
    run_name.bold = True
    run_name.font.name = 'Cambria'
    run_name.font.size = Pt(9)

    # Consortium (from linked project or fallback to record.consortium_info)
    consortium_text = _build_consortium_text(record, cp)
    if consortium_text:
        cell.add_paragraph()

        p_ci_label = cell.add_paragraph()
        run_ci = p_ci_label.add_run('Consortium')
        run_ci.bold = True
        run_ci.font.name = 'Cambria'
        run_ci.font.size = Pt(9)

        p_ci_text = cell.add_paragraph()
        run_ci_text = p_ci_text.add_run(consortium_text)
        run_ci_text.font.name = 'Cambria'
        run_ci_text.font.size = Pt(9)

    # Brief Description
    if record.description:
        # 빈 줄
        cell.add_paragraph()

        p_bd_label = cell.add_paragraph()
        run_bd = p_bd_label.add_run('Brief Description')
        run_bd.bold = True
        run_bd.font.name = 'Cambria'
        run_bd.font.size = Pt(9)

        p_bd_text = cell.add_paragraph()
        run_bd_text = p_bd_text.add_run(record.description)
        run_bd_text.font.name = 'Cambria'
        run_bd_text.font.size = Pt(9)

    # Deliverables / Outputs
    if record.achievements:
        p_do_label = cell.add_paragraph()
        run_do = p_do_label.add_run('Deliverables / Outputs')
        run_do.bold = True
        run_do.font.name = 'Cambria'
        run_do.font.size = Pt(9)

        p_do_text = cell.add_paragraph()
        run_do_text = p_do_text.add_run(record.achievements)
        run_do_text.font.name = 'Cambria'
        run_do_text.font.size = Pt(9)

    # 모든 paragraph의 간격 최소화
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)


def _format_usd_million(value):
    """USD 백만 단위 포맷팅"""
    if value and float(value) > 0:
        return f"US$ {float(value):.3f} mil."
    return '-'
