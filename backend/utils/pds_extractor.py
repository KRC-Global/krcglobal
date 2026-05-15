"""
GBMS - PDS 데이터 자동 추출 유틸리티 (외부망 / Cloudflare R2)
업로드된 PDF/DOCX 문서에서 PDS(Project Data Sheet) 양식 항목을 정규식으로 추출한다.

내부망 버전과 달리 첨부 파일은 로컬 disk가 아니라 R2 버킷에 저장되어 있으므로
R2 → BytesIO 로 받아 메모리 내에서 텍스트 추출한다.
"""
import os
import re
from datetime import datetime
from io import BytesIO

# 파일 크기 & 페이지 제한
MAX_FILE_SIZE_MB = 20
MAX_PDF_PAGES = 50


def _download_r2_bytes(key, max_bytes):
    """R2에서 객체를 다운로드. 크기 초과면 None."""
    if not key:
        return None
    try:
        from utils.r2_storage import download_file
        obj = download_file(key)
        size = obj.get('ContentLength') or 0
        if max_bytes and size and size > max_bytes:
            return None
        data = obj['Body'].read()
        if max_bytes and len(data) > max_bytes:
            return None
        return data
    except Exception as e:
        print(f'[PDS Extractor] R2 다운로드 실패 ({key}): {e}')
        return None


def _extract_text_from_pdf_bytes(raw):
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(BytesIO(raw)) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= MAX_PDF_PAGES:
                    break
                t = page.extract_text()
                if t:
                    texts.append(t)
        return '\n'.join(texts)
    except Exception as e:
        print(f'[PDS Extractor] PDF 추출 실패: {e}')
        return ''


def _extract_text_from_docx_bytes(raw):
    try:
        from docx import Document
        doc = Document(BytesIO(raw))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())
        return '\n'.join(texts)
    except Exception as e:
        print(f'[PDS Extractor] DOCX 추출 실패: {e}')
        return ''


def _extract_text_from_r2(key, original_name=None):
    """R2 키에서 텍스트 추출. 확장자는 original_name 또는 key 에서 추론."""
    raw = _download_r2_bytes(key, MAX_FILE_SIZE_MB * 1024 * 1024)
    if not raw:
        return ''
    name_for_ext = (original_name or key or '').lower()
    if name_for_ext.endswith('.pdf'):
        return _extract_text_from_pdf_bytes(raw)
    if name_for_ext.endswith('.docx'):
        return _extract_text_from_docx_bytes(raw)
    return ''


def _parse_address(text):
    m = re.search(r'[Aa]ddress\s*[:\-]\s*(.+?)(?:\n|$)', text)
    if m:
        addr = m.group(1).strip()
        if len(addr) > 5:
            return addr
    return None


def _parse_staff_months(text):
    m = re.search(r'(\d[\d,]*\.?\d*)\s*(?:staff[\s\-]?months?|person[\s\-]?months?)', text, re.I)
    if m:
        return m.group(0).strip()
    return None


def _parse_senior_staff(text):
    patterns = [
        r'(?:Team\s+Leader|Project\s+(?:Manager|Director)|Chief\s+(?:Engineer|Consultant))\s*[:\-]?\s*([^\n]{3,80})',
    ]
    results = []
    for pat in patterns:
        for m in re.finditer(pat, text, re.I):
            results.append(m.group(0).strip())
    if results:
        return '; '.join(results[:10])
    return None


def _parse_location(text):
    patterns = [
        r'(?:Province|District|Region|Municipality|City|State)\s*(?:of\s+)?[:\-]?\s*([A-Z][a-zA-Z\s,]{2,60})',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return m.group(0).strip()
    return None


def _parse_narrative(text):
    patterns = [
        r'(?:Narrative\s+[Dd]escription|Project\s+[Dd]escription|Description\s+of\s+(?:the\s+)?[Pp]roject)\s*[:\-]?\s*(.{30,1500}?)(?:\n\n|\n[A-Z0-9])',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.DOTALL)
        if m:
            return m.group(1).strip()
    return None


def _parse_services(text):
    patterns = [
        r'(?:Scope\s+of\s+(?:[Ss]ervices|[Ww]ork)|Description\s+of\s+(?:actual\s+)?[Ss]ervices|Services\s+[Pp]rovided)\s*[:\-]?\s*(.{30,1500}?)(?:\n\n|\n[A-Z0-9])',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.DOTALL)
        if m:
            return m.group(1).strip()
    return None


def _gather_project_files(project_id):
    """사업에 연결된 모든 첨부의 R2 키와 표시명을 우선순위 순으로 수집.

    반환: ([(r2_key, source_label, original_name), ...], skipped_names)
    """
    from models import Contract, Proposal, PerformanceRecord, TorRfp, Eoi

    files = []
    skipped = []

    # Contract — 최우선
    try:
        for c in Contract.query.filter_by(consulting_project_id=project_id).all():
            if c.file_path:
                files.append((c.file_path, f'Contract:{c.file_name or c.file_path}', c.file_name))
    except Exception:
        pass

    # Proposal — 기술제안서 우선, 그 다음 단일 파일 호환 컬럼
    try:
        for p in Proposal.query.filter_by(consulting_project_id=project_id).all():
            for path_attr, name_attr in [
                ('technical_file_path', 'technical_file_name'),
                ('file_path', 'file_name'),
            ]:
                fp = getattr(p, path_attr, None)
                fn = getattr(p, name_attr, None)
                if fp:
                    files.append((fp, f'Proposal:{fn or fp}', fn))
    except Exception:
        pass

    # PerformanceRecord
    try:
        for pr in PerformanceRecord.query.filter_by(consulting_project_id=project_id).all():
            if pr.file_path:
                files.append((pr.file_path, f'Performance:{pr.file_name or pr.file_path}', pr.file_name))
    except Exception:
        pass

    # TorRfp
    try:
        for tr in TorRfp.query.filter_by(consulting_project_id=project_id).all():
            for path_attr, name_attr in [
                ('tor_file_path', 'tor_file_name'),
                ('rfp_file_path', 'rfp_file_name'),
            ]:
                fp = getattr(tr, path_attr, None)
                fn = getattr(tr, name_attr, None)
                if fp:
                    files.append((fp, f'TorRfp:{fn or fp}', fn))
    except Exception:
        pass

    # Eoi
    try:
        for e in Eoi.query.filter_by(consulting_project_id=project_id).all():
            if e.eoi_file_path:
                files.append((e.eoi_file_path,
                              f'Eoi:{e.eoi_file_name or e.eoi_file_path}', e.eoi_file_name))
    except Exception:
        pass

    return files, skipped


def extract_pds_data(project_id):
    """사업의 모든 첨부 문서를 R2에서 받아 텍스트화하고 PDS 필드를 정규식으로 추출.

    Returns:
        dict {
            'fields':       { field_name: value, ... },
            'updatedFields':[field_name, ...],
            'skippedFiles': [name, ...],
            'warnings':     [str, ...],
            'processedFiles': int,
        }
    """
    from models import db, ConsultingProject

    project = ConsultingProject.query.get(project_id)
    if not project:
        return {'fields': {}, 'updatedFields': [], 'skippedFiles': [],
                'warnings': ['프로젝트를 찾을 수 없습니다.'], 'processedFiles': 0}

    file_list, skipped = _gather_project_files(project_id)
    warnings = []
    if not file_list:
        warnings.append('분석할 첨부 파일이 없습니다.')
        return {'fields': {}, 'updatedFields': [], 'skippedFiles': skipped,
                'warnings': warnings, 'processedFiles': 0}

    all_text = ''
    processed = 0
    for r2_key, source, original_name in file_list:
        text = _extract_text_from_r2(r2_key, original_name)
        if text:
            all_text += f'\n--- {source} ---\n{text}\n'
            processed += 1
        else:
            skipped.append(source)

    if not all_text:
        warnings.append('첨부 파일에서 텍스트를 추출하지 못했습니다.')
        return {'fields': {}, 'updatedFields': [], 'skippedFiles': skipped,
                'warnings': warnings, 'processedFiles': 0}

    parsers = {
        'pds_location_within_country': _parse_location,
        'pds_client_address': _parse_address,
        'pds_total_staff_months': _parse_staff_months,
        'pds_senior_staff': _parse_senior_staff,
        'pds_narrative_description_en': _parse_narrative,
        'pds_services_description_en': _parse_services,
    }

    extracted = {}
    updated_fields = []

    for field_name, parser_fn in parsers.items():
        current = getattr(project, field_name, None)
        if current:
            continue  # 사용자가 이미 채운 값은 보호
        value = parser_fn(all_text)
        if value:
            extracted[field_name] = value
            setattr(project, field_name, value)
            updated_fields.append(field_name)

    if updated_fields:
        project.pds_extracted_at = datetime.utcnow()
        db.session.commit()

    return {
        'fields': extracted,
        'updatedFields': updated_fields,
        'skippedFiles': skipped,
        'warnings': warnings,
        'processedFiles': processed,
    }
