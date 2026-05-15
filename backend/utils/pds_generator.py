"""
GBMS - PDS(Project Data Sheet) DOCX 생성 유틸리티
(ex)PDS.docx 템플릿 기반으로 8행 2열 PDS 표를 생성한다.
"""
import os
import re
import copy
from io import BytesIO
from datetime import datetime


def _format_date_pds(date_str):
    """'YY.MM 또는 YYYY.MM 형식을 MM/YYYY로 변환 (PDS 양식용)"""
    if not date_str:
        return '-'
    clean = str(date_str).replace("'", '').replace('"', '').replace('\u2018', '').replace('\u2019', '').strip()
    parts = re.split(r'[-./]', clean)
    if len(parts) >= 2:
        yy = int(parts[0])
        mm = int(parts[1])
        if yy < 100:
            year = 1900 + yy if yy >= 50 else 2000 + yy
        else:
            year = yy
        return f'{mm:02d}/{year}'
    return clean


def _calc_duration_months(start_str, end_str):
    """두 날짜 문자열 사이의 개월 수 계산"""
    try:
        def parse_ym(s):
            if not s:
                return None
            clean = str(s).replace("'", '').replace('"', '').replace('\u2018', '').replace('\u2019', '').strip()
            parts = re.split(r'[-./]', clean)
            if len(parts) >= 2:
                yy = int(parts[0])
                mm = int(parts[1])
                if yy < 100:
                    year = 1900 + yy if yy >= 50 else 2000 + yy
                else:
                    year = yy
                return year, mm
            return None

        s = parse_ym(start_str)
        e = parse_ym(end_str)
        if s and e:
            months = (e[0] - s[0]) * 12 + (e[1] - s[1]) + 1
            return max(months, 1)
    except Exception:
        pass
    return None


def _format_currency(value, unit='백만원'):
    """숫자를 통화 형식으로 포맷"""
    if value is None or value == 0:
        return '-'
    try:
        val = float(value)
        if val >= 1:
            return f'{val:,.2f} {unit}'
        else:
            return f'{val:,.4f} {unit}'
    except (ValueError, TypeError):
        return '-'


def _format_usd(value):
    """USD 백만 단위 포맷"""
    if value is None or value == 0:
        return None
    try:
        val = float(value)
        return f'USD {val:,.2f} million'
    except (ValueError, TypeError):
        return None


def _build_consortium_names(project):
    """컨소시엄 구성원 목록 생성"""
    names = []
    if project.lead_company:
        ratio = f' ({float(project.lead_company_ratio)*100:.0f}%)' if project.lead_company_ratio else ''
        names.append(f'{project.lead_company}{ratio} (Lead)')

    for i in range(1, 6):
        jv = getattr(project, f'jv{i}', None)
        if jv:
            ratio = getattr(project, f'jv{i}_ratio', None)
            ratio_str = f' ({float(ratio)*100:.0f}%)' if ratio else ''
            names.append(f'{jv}{ratio_str}')

    return '\n'.join(names) if names else '-'


def build_pds_dict(project):
    """8행 2열 PDS 표의 16개 셀 텍스트를 구성한다.

    Returns:
        list of 8 tuples: [(left_text, right_text), ...]
        행 7, 8은 right_text가 None (병합 행)
    """
    # Row 1: Assignment name / Approximate value
    budget_text = '-'
    usd = _format_usd(project.budget_usd)
    if usd:
        budget_text = usd
    elif project.budget:
        budget_text = _format_currency(float(project.budget))

    row1_left = f"Assignment name:\n{project.title_en or project.title_kr or '-'}"
    row1_right = f"Approximate value of the contract:\n{budget_text}"

    # Row 2: Country + Location / Duration
    location = project.pds_location_within_country or '-'
    country = project.country or '-'
    duration = _calc_duration_months(project.start_date, project.end_date)
    start_fmt = _format_date_pds(project.start_date)
    end_fmt = _format_date_pds(project.end_date)
    duration_text = f"{duration} months ({start_fmt} ~ {end_fmt})" if duration else f"{start_fmt} ~ {end_fmt}"

    row2_left = f"Country: {country}\nLocation within country: {location}"
    row2_right = f"Duration of assignment (months):\n{duration_text}"

    # Row 3: Client / Total staff-months
    row3_left = f"Name of Client:\n{project.client or '-'}"
    row3_right = f"Total No. of staff-months of the assignment:\n{project.pds_total_staff_months or '-'}"

    # Row 4: Address / Approximate value of KRC services
    krc_value = '-'
    krc_usd = _format_usd(project.krc_budget_usd)
    if krc_usd:
        krc_value = krc_usd
    elif project.krc_budget:
        krc_value = _format_currency(float(project.krc_budget))
    if project.krc_share_ratio:
        krc_value += f' ({float(project.krc_share_ratio)*100:.1f}%)'

    row4_left = f"Address:\n{project.pds_client_address or '-'}"
    row4_right = f"Approximate value of the services provided by your firm under the contract:\n{krc_value}"

    # Row 5: Start/Completion / Associated staff-months
    row5_left = f"Start date (Month/Year): {start_fmt}\nCompletion date (Month/Year): {end_fmt}"
    row5_right = f"No. of professional staff-months provided by associated Consultants:\n{project.pds_associated_staff_months or '-'}"

    # Row 6: Associated Consultants / Senior staff
    row6_left = f"Name of associated Consultants, if any:\n{_build_consortium_names(project)}"
    row6_right = f"Name of senior professional staff of your firm involved and functions performed:\n{project.pds_senior_staff or '-'}"

    # Row 7: Narrative description (merged)
    narrative = project.pds_narrative_description_en or project.description_en or project.description or '-'
    row7_left = f"Narrative description of Project:\n\n{narrative}"

    # Row 8: Description of actual services (merged)
    services = project.pds_services_description_en or '-'
    row8_left = f"Description of actual services provided by your staff within the assignment:\n\n{services}"

    return [
        (row1_left, row1_right),
        (row2_left, row2_right),
        (row3_left, row3_right),
        (row4_left, row4_right),
        (row5_left, row5_right),
        (row6_left, row6_right),
        (row7_left, None),
        (row8_left, None),
    ]


def _set_cell_text(cell, text):
    """셀 텍스트를 설정하되 기존 서식을 최대한 유지"""
    # 기존 paragraph의 run 서식을 참조
    if cell.paragraphs:
        first_para = cell.paragraphs[0]
        # 기존 run이 있으면 서식 정보 보존
        ref_run = first_para.runs[0] if first_para.runs else None

        # 기존 paragraph 모두 제거 (첫 번째 제외)
        for para in cell.paragraphs[1:]:
            p_element = para._element
            p_element.getparent().remove(p_element)

        # 첫 번째 paragraph의 텍스트만 교체
        first_para.clear()

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if i == 0:
                run = first_para.add_run(line)
            else:
                # 줄바꿈은 새 paragraph 추가
                from docx.oxml.ns import qn
                new_para = copy.deepcopy(first_para._element)
                # 새 paragraph의 run 텍스트만 교체
                for r in new_para.findall(qn('w:r')):
                    new_para.remove(r)
                cell._element.append(new_para)
                from docx.text.paragraph import Paragraph
                new_p = Paragraph(new_para, cell)
                run = new_p.add_run(line)

            # 참조 서식 적용
            if ref_run and ref_run.font:
                if ref_run.font.name:
                    run.font.name = ref_run.font.name
                if ref_run.font.size:
                    run.font.size = ref_run.font.size


def _keep_first_table_only(doc, table):
    """Keep the first PDS table and remove the remaining sample tables."""
    body = doc.element.body
    sectPr_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
    table_elem = table._element
    found_first_table = False

    for child in list(body):
        if child is table_elem or child == table_elem:
            found_first_table = True
            continue

        if not found_first_table:
            continue

        if child.tag == sectPr_tag:
            continue

        body.remove(child)


def generate_pds_docx(project):
    """PDS DOCX 파일을 생성하여 BytesIO로 반환한다.

    템플릿 파일이 있으면 첫 번째 테이블의 셀 텍스트를 치환하고,
    없으면 fallback으로 새 테이블을 생성한다.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pds_data = build_pds_dict(project)

    template_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'templates', 'pds_template.docx'
    )

    if os.path.exists(template_path):
        # 템플릿 기반 생성
        doc = Document(template_path)

        if doc.tables:
            table = doc.tables[0]

            _keep_first_table_only(doc, table)

            # 8행 셀 텍스트 치환
            for row_idx, (left_text, right_text) in enumerate(pds_data):
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    # 좌측 셀
                    _set_cell_text(row.cells[0], left_text)
                    # 우측 셀 (병합 행이 아닌 경우)
                    if right_text is not None and len(row.cells) > 1:
                        _set_cell_text(row.cells[1], right_text)
    else:
        # Fallback: 새 문서 생성
        doc = Document()

        # 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run('Project Data Sheet (PDS) for consultant\'s experience')
        run.bold = True
        run.font.size = Pt(12)

        # 8행 2열 테이블
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        for row_idx, (left_text, right_text) in enumerate(pds_data):
            row = table.rows[row_idx]
            row.cells[0].text = left_text

            if right_text is not None:
                row.cells[1].text = right_text
            else:
                # 병합 행
                row.cells[0].merge(row.cells[1])

    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
